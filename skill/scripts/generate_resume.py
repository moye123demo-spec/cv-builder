#!/usr/bin/env python3
"""Generate editable resumes from verified JSON and JSON-driven templates."""

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from export_pdf import export_pdf
from validate_profile import load_profile, validate


STYLE_DIR = Path(__file__).resolve().parents[1] / "assets" / "styles"

LABELS = {
    "cn": {
        "summary": "职业概述", "education": "教育背景", "skills": "技术能力",
        "experience": "工作经历", "projects": "代表项目", "achievements": "成果与补充信息",
        "research": "研究方向", "publications": "论文与学术成果", "funding": "科研项目与经费",
        "teaching": "教学与学术服务", "photo": "照片\n可选",
    },
    "en": {
        "summary": "Professional Summary", "education": "Education", "skills": "Technical Skills",
        "experience": "Experience", "projects": "Selected Projects", "achievements": "Achievements",
        "research": "Research Interests", "publications": "Publications & Academic Output",
        "funding": "Funding & Research Projects", "teaching": "Teaching & Academic Service",
        "photo": "Optional\nPhoto",
    },
}


def available_templates() -> list[str]:
    return sorted(json.loads(path.read_text(encoding="utf-8"))["template"] for path in STYLE_DIR.glob("*.json"))


def load_template(name: str) -> dict:
    path = STYLE_DIR / f"{name.replace('-', '_')}.json"
    if not path.exists():
        raise ValueError(f"Unknown template: {name}")
    template = json.loads(path.read_text(encoding="utf-8"))
    required = {"template", "accent_color", "title_color", "language", "photo_optional", "bilingual_headings"}
    missing = sorted(required - set(template))
    if missing:
        raise ValueError(f"Template {name} is missing: {', '.join(missing)}")
    if template["template"] != name:
        raise ValueError(f"Template name mismatch in {path.name}")
    return template


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_font(run, size: float, template: dict, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = template.get("font_en", "Aptos")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), template.get("font_cn", "Microsoft YaHei"))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_cell_width(cell, width_cm: float) -> None:
    props = cell._tc.get_or_add_tcPr()
    width = OxmlElement("w:tcW")
    width.set(qn("w:w"), str(int(width_cm * 567)))
    width.set(qn("w:type"), "dxa")
    props.append(width)


def heading_text(key: str, template: dict) -> str:
    base = LABELS[template["language"]][key]
    if template["bilingual_headings"]:
        return f"{base}  ({LABELS['en'][key]})"
    return base


def add_section(doc: Document, key: str, template: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(heading_text(key, template)), 11.2, template, bold=True, color=rgb(template["accent_color"]))


def add_bullets(doc: Document, bullets: list[str], template: dict) -> None:
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_together = True
        set_font(p.add_run(item), 9, template)


def add_role(doc: Document, role: dict, primary: RGBColor, template: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(f"{role['title']}  |  {role['institution']}"), 9.5, template, bold=True, color=primary)
    set_font(p.add_run(f"    {role['dates']}"), 9, template, bold=True)
    add_bullets(doc, role.get("bullets", []), template)


def add_record_list(doc: Document, records, primary: RGBColor, template: dict) -> None:
    for record in records:
        if isinstance(record, str):
            add_bullets(doc, [record], template)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        set_font(p.add_run(record.get("title") or record.get("name") or "Verified record"), 9.5, template, bold=True, color=primary)
        if record.get("details"):
            set_font(p.add_run(f"  |  {record['details']}"), 9, template)
        add_bullets(doc, record.get("bullets", []), template)


def add_header(doc: Document, profile: dict, template: dict, photo: Path | None) -> None:
    primary = rgb(template["title_color"])
    accent = rgb(template["accent_color"])
    include_photo = template["photo_optional"]
    header = doc.add_table(rows=1, cols=2 if include_photo else 1)
    header.autofit = False
    left = header.rows[0].cells[0]
    set_cell_width(left, 13.7 if include_photo else 16.7)
    p = left.paragraphs[0]
    set_font(p.add_run(profile["name"]), 23, template, bold=True, color=primary)
    p = left.add_paragraph()
    set_font(p.add_run(profile["headline"]), 10.8, template, bold=True, color=accent)
    contact = profile["contact"]
    fields = [contact.get("phone"), contact.get("email"), profile.get("location"), contact.get("github")]
    p = left.add_paragraph()
    set_font(p.add_run(" | ".join(str(x) for x in fields if x)), 8.6, template)
    if not include_photo:
        return
    right = header.rows[0].cells[1]
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_width(right, 3.0)
    if photo:
        if not photo.exists():
            raise ValueError(f"Photo not found: {photo}")
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        right.paragraphs[0].add_run().add_picture(str(photo), width=Cm(2.45))
    else:
        shade(right, "F2F4F5")
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(LABELS[template["language"]]["photo"]), 8, template, color=primary)


def build(profile: dict, output: Path, template_name: str, photo: Path | None) -> None:
    template = load_template(template_name)
    primary = rgb(template["title_color"])
    doc = Document()
    page = doc.sections[0]
    page.top_margin = Cm(template.get("top_margin_cm", 1.25))
    page.bottom_margin = Cm(template.get("bottom_margin_cm", 1.25))
    page.left_margin = Cm(template.get("left_margin_cm", 1.45))
    page.right_margin = Cm(template.get("right_margin_cm", 1.45))

    add_header(doc, profile, template, photo)
    add_section(doc, "summary", template)
    p = doc.add_paragraph()
    set_font(p.add_run(profile["summary"]), 9, template)
    add_section(doc, "education", template)
    for item in profile["education"]:
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = True
        set_font(p.add_run(f"{item['degree']}  |  {item['institution']}  |  {item['dates']}"), 9, template)
        for note in item.get("notes", []):
            add_bullets(doc, [note], template)
    add_section(doc, "skills", template)
    for label, values in profile["skills"].items():
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = True
        set_font(p.add_run(label.replace("_", " ").title() + ": "), 9, template, bold=True, color=primary)
        set_font(p.add_run("；".join(values) if template["language"] == "cn" else "; ".join(values)), 9, template)

    if template_name.startswith("academic"):
        for key, records in [("research", profile.get("research_interests", [])), ("publications", profile.get("publications", [])), ("funding", profile.get("funding", [])), ("teaching", profile.get("teaching_service", []))]:
            if records:
                add_section(doc, key, template)
                add_record_list(doc, records, primary, template)

    add_section(doc, "experience", template)
    for role in profile["experience"]:
        add_role(doc, role, primary, template)
    if profile.get("projects"):
        add_section(doc, "projects", template)
        add_record_list(doc, profile["projects"], primary, template)
    if profile.get("achievements"):
        add_section(doc, "achievements", template)
        if template["language"] == "en" and len(profile["achievements"]) <= 3:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            set_font(p.add_run("  •  ".join(profile["achievements"])), 8.6, template)
        else:
            add_bullets(doc, profile["achievements"], template)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("profile", type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--template", choices=available_templates(), default="industry-cn")
    parser.add_argument("--photo", type=Path)
    parser.add_argument("--pdf", action="store_true")
    parser.add_argument("--pdf-engine", choices=("auto", "libreoffice", "word"), default="auto")
    args = parser.parse_args()
    profile = load_profile(args.profile)
    errors = validate(profile)
    if errors:
        print("\n".join(f"ERROR: {item}" for item in errors), file=sys.stderr)
        return 1
    build(profile, args.output, args.template, args.photo)
    print(f"Created DOCX: {args.output} ({args.template})")
    if args.pdf:
        print(f"Created PDF: {export_pdf(args.output, args.pdf_engine)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
